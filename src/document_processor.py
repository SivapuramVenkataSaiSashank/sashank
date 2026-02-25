"""
document_processor.py
Handles text extraction from PDF, DOCX, and EPUB files.
"""
import os
import re
import uuid
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

class DocumentProcessor:
    def __init__(self):
        self.pages = []          # list of {"index": int, "text": str, "label": str, "file_id": int}
        self.loaded_files = []   # list of {"id": int, "path": str, "name": str, "type": str, "page_count": int}
        self.next_file_id = 1
        
        self.current_page = 0
        self.file_path = None
        self.doc_type = None
        self.title = "Untitled Document"
        
        # ChromaDB setup
        self.chroma_client = chromadb.Client(Settings(anonymized_telemetry=False))
        # Use simple built-in default embeddings (all-MiniLM-L6-v2) for speed
        self.embedding_fn = embedding_functions.DefaultEmbeddingFunction()
        self.collection_name = "doc_collection"
        self.collection = None

    # ─────────────────────────── Public API ───────────────────────────

    def load(self, filepath: str, append: bool = False) -> bool:
        """Load a document and return True on success. If append is True, add to existing pages."""
        ext = os.path.splitext(filepath)[1].lower()
        new_title = os.path.basename(filepath)
        
        if not append:
            self.pages = []
            self.loaded_files = []
            self.current_page = 0
            self.file_path = filepath
            self.title = new_title
        else:
            if self.title != "Untitled Document" and len(self.loaded_files) > 0:
                # Instead of appending indefinitely, "Multiple Files (N)" is cleaner if > 1
                self.title = f"Multiple Files ({len(self.loaded_files) + 1})"
            else:
                self.title = new_title

        current_file_id = self.next_file_id
        self.next_file_id += 1
        start_pages_len = len(self.pages)

        try:
            if ext == ".pdf":
                self._load_pdf(filepath, current_file_id)
                self.doc_type = "PDF" if not append else "Mixed"
            elif ext in (".docx", ".doc"):
                self._load_docx(filepath, current_file_id)
                self.doc_type = "DOCX" if not append else "Mixed"
            elif ext == ".epub":
                self._load_epub(filepath, current_file_id)
                self.doc_type = "EPUB" if not append else "Mixed"
            elif ext == ".txt":
                self._load_txt(filepath, current_file_id)
                self.doc_type = "TXT" if not append else "Mixed"
            else:
                return False
                
            # Record file metadata
            pages_added = len(self.pages) - start_pages_len
            self.loaded_files.append({
                "id": current_file_id,
                "path": filepath,
                "name": new_title,
                "type": ext.upper().replace('.', ''),
                "page_count": pages_added
            })
                
            if len(self.pages) > 0:
                self._build_vector_index()
                return True
            return False
        except Exception as e:
            print(f"[DocumentProcessor] Load error: {e}")
            return False

    def unload(self, delete_file: bool = False) -> bool:
        """Clear document from memory and optionally delete the file from disk."""
        try:
            if delete_file:
                for f in self.loaded_files:
                    path = f["path"]
                    if path and os.path.isfile(path):
                        try: os.remove(path)
                        except Exception as e: print(f"Could not delete {path}: {e}")

            if self.collection:
                try:
                    self.chroma_client.delete_collection(name=self.collection_name)
                    self.collection = None
                except Exception:
                    pass

            self.pages = []
            self.loaded_files = []
            self.current_page = 0
            self.file_path = None
            self.doc_type = None
            self.title = "Untitled Document"
            return True
        except Exception as e:
            print(f"[DocumentProcessor] Unload error: {e}")
            return False

    def remove_file(self, file_id: int, delete_file: bool = True) -> bool:
        """Remove a specific file from the loaded batch."""
        target_file = next((f for f in self.loaded_files if f["id"] == file_id), None)
        if not target_file: return False

        # If it's the last file, just full unload
        if len(self.loaded_files) == 1:
            return self.unload(delete_file=delete_file)

        # Remove physical file
        if delete_file and target_file["path"] and os.path.isfile(target_file["path"]):
            try: os.remove(target_file["path"])
            except Exception as e: print(f"Could not delete {target_file['path']}: {e}")

        # Remove from state
        self.loaded_files = [f for f in self.loaded_files if f["id"] != file_id]
        
        # Filter pages
        self.pages = [p for p in self.pages if p.get("file_id") != file_id]
        
        # Reset current page if out of bounds
        if self.current_page >= len(self.pages):
            self.current_page = max(0, len(self.pages) - 1)
            
        # Re-index remaining pages
        for i, p in enumerate(self.pages):
            p["index"] = i

        # Rebuild title
        if len(self.loaded_files) == 1:
            self.title = self.loaded_files[0]["name"]
            self.doc_type = self.loaded_files[0]["type"]
        else:
            self.title = f"Multiple Files ({len(self.loaded_files)})"
            self.doc_type = "Mixed"
            
        # Rebuild vector index
        self._build_vector_index()
        return True

    def page_count(self) -> int:
        return len(self.pages)

    def get_page(self, index: int) -> str:
        if 0 <= index < len(self.pages):
            return self.pages[index]["text"]
        return ""

    def get_current_text(self) -> str:
        return self.get_page(self.current_page)

    def get_current_label(self) -> str:
        if 0 <= self.current_page < len(self.pages):
            return self.pages[self.current_page]["label"]
        return "Unknown"

    def next_page(self) -> bool:
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            return True
        return False

    def prev_page(self) -> bool:
        if self.current_page > 0:
            self.current_page -= 1
            return True
        return False

    def go_to_page(self, index: int) -> bool:
        if 0 <= index < len(self.pages):
            self.current_page = index
            return True
        return False

    def get_full_text(self, max_chars: int = 50000) -> str:
        """Return combined text of all pages (truncated for AI)."""
        combined = "\n\n".join(p["text"] for p in self.pages)
        return combined[:max_chars]

    def get_chapter_text(self, chapter_num: int) -> str:
        """Return text of a specific chapter/page (1-indexed)."""
        idx = chapter_num - 1
        if 0 <= idx < len(self.pages):
            return self.pages[idx]["text"]
        return ""

    def search(self, query: str) -> list:
        """Search text and return list of (page_index, snippet) matches."""
        results = []
        q = query.lower()
        for page in self.pages:
            text = page["text"].lower()
            pos = text.find(q)
            if pos != -1:
                snippet = page["text"][max(0, pos-60):pos+120]
                results.append({"page": page["index"], "label": page["label"], "snippet": snippet})
        return results

    def get_relevant_context(self, query: str, n_results: int = 4) -> str:
        """Fetch the most relevant text chunks for a given query via ChromaDB."""
        if not self.collection or self.collection.count() == 0:
            return self.get_full_text(max_chars=10000)

        # Ensure we don't ask for more results than we have chunks
        k = min(n_results, self.collection.count())
        if k == 0:
            return ""
            
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=k
            )
            
            if not results["documents"] or not results["documents"][0]:
                return self.get_full_text(max_chars=10000)
                
            # Combine the returned chunks
            documents = results["documents"][0]
            context = "\n...\n".join(documents)
            return context
        except Exception as e:
            print(f"[DocumentProcessor] Vector search failed: {e}. Falling back to full text.")
            return self.get_full_text(max_chars=10000)

    # ─────────────────────────── Private loaders ───────────────────────

    def _build_vector_index(self):
        """Index all loaded pages into an ephemeral Chroma collection."""
        try:
            # Recreate collection to clear old data
            try:
                self.chroma_client.delete_collection(name=self.collection_name)
            except Exception:
                pass
                
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name, 
                embedding_function=self.embedding_fn
            )
            
            # Simple chunking logic to ensure we don't hit max payload sizes
            docs = []
            metadatas = []
            ids = []
            
            for page in self.pages:
                text = page["text"]
                # Chunk aggressively for better vector retrieval
                words = text.split()
                chunk_size = 300
                overlap = 50
                
                if not words: continue
                
                for i in range(0, len(words), chunk_size - overlap):
                    chunk = " ".join(words[i:i + chunk_size])
                    docs.append(chunk)
                    metadatas.append({"page": page["index"], "label": page["label"]})
                    ids.append(str(uuid.uuid4()))
                    
            if docs:
                self.collection.add(
                    documents=docs,
                    metadatas=metadatas,
                    ids=ids
                )
                print(f"[DocumentProcessor] Indexed {len(docs)} chunks into ChromaDB.")
        except Exception as e:
            print(f"[DocumentProcessor] Error bulding vector index: {e}")
            self.collection = None

    def _load_pdf(self, filepath: str, file_id: int):
        import fitz  # PyMuPDF
        start_idx = len(self.pages)
        doc = fitz.open(filepath)
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                self.pages.append({
                    "index": start_idx + i,
                    "text": text,
                    "label": f"Page {start_idx + i + 1}",
                    "file_id": file_id
                })
        doc.close()

    def _load_docx(self, filepath: str, file_id: int):
        from docx import Document
        doc = Document(filepath)
        start_idx = len(self.pages)
        # Split by headings into chapters, otherwise into chunks
        current_chunk = []
        chapter_idx = start_idx
        chapter_label = f"Section {start_idx + 1}"

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                if current_chunk:
                    self.pages.append({
                        "index": chapter_idx,
                        "text": "\n".join(current_chunk),
                        "label": chapter_label,
                        "file_id": file_id
                    })
                    chapter_idx += 1
                chapter_label = para.text.strip() or f"Section {chapter_idx + 1}"
                current_chunk = []
            else:
                if para.text.strip():
                    current_chunk.append(para.text.strip())

        if current_chunk:
            self.pages.append({
                "index": chapter_idx,
                "text": "\n".join(current_chunk),
                "label": chapter_label,
                "file_id": file_id
            })

        # If no headings found, chunk by ~500 words
        if not self.pages or all(p.get("file_id") != file_id for p in self.pages):
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            words = all_text.split()
            chunk_size = 500
            for i in range(0, len(words), chunk_size):
                chunk = " ".join(words[i:i + chunk_size])
                self.pages.append({
                    "index": start_idx + (i // chunk_size),
                    "text": chunk,
                    "label": f"Section {start_idx + (i // chunk_size) + 1}",
                    "file_id": file_id
                })

    def _load_epub(self, filepath: str, file_id: int):
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(filepath)
        start_idx = len(self.pages)
        chapter_idx = start_idx
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                text = soup.get_text(separator="\n").strip()
                if len(text) > 100:
                    title_tag = soup.find(["h1", "h2", "h3"])
                    label = title_tag.get_text().strip() if title_tag else f"Chapter {chapter_idx + 1}"
                    self.pages.append({
                        "index": chapter_idx,
                        "text": text,
                        "label": label,
                        "file_id": file_id
                    })
                    chapter_idx += 1

    def _load_txt(self, filepath: str, file_id: int):
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        words = content.split()
        chunk_size = 600
        start_idx = len(self.pages)
        for i in range(0, len(words), chunk_size):
            chunk = " ".join(words[i:i + chunk_size])
            self.pages.append({
                "index": start_idx + (i // chunk_size),
                "text": chunk,
                "label": f"Section {start_idx + (i // chunk_size) + 1}",
                "file_id": file_id
            })
